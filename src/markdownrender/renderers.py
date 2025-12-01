"""
Renderers for different output formats: HTML, PDF, Word (DOCX), Excel (XLSX).
"""

import io
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment

from .parser import MarkdownParser


class HTMLRenderer:
    """Render Markdown to HTML with styling."""

    DEFAULT_CSS = """
    <style>
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }
        h1, h2, h3, h4, h5, h6 {
            margin-top: 1.5em;
            margin-bottom: 0.5em;
            font-weight: 600;
        }
        h1 { font-size: 2em; border-bottom: 1px solid #eee; padding-bottom: 0.3em; }
        h2 { font-size: 1.5em; border-bottom: 1px solid #eee; padding-bottom: 0.3em; }
        h3 { font-size: 1.25em; }
        code {
            background-color: #f4f4f4;
            padding: 0.2em 0.4em;
            border-radius: 3px;
            font-family: 'SF Mono', Consolas, monospace;
            font-size: 0.9em;
        }
        pre {
            background-color: #f4f4f4;
            padding: 16px;
            border-radius: 6px;
            overflow-x: auto;
        }
        pre code {
            background: none;
            padding: 0;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }
        th, td {
            border: 1px solid #ddd;
            padding: 8px 12px;
            text-align: left;
        }
        th {
            background-color: #f4f4f4;
            font-weight: 600;
        }
        tr:nth-child(even) {
            background-color: #f9f9f9;
        }
        blockquote {
            margin: 1em 0;
            padding: 0.5em 1em;
            border-left: 4px solid #ddd;
            color: #666;
            background-color: #f9f9f9;
        }
        img {
            max-width: 100%;
            height: auto;
        }
        .mermaid-diagram, .plantuml-diagram {
            margin: 1em 0;
            text-align: center;
        }
        .mermaid {
            background-color: white;
        }
        .highlight {
            background-color: #f4f4f4;
            padding: 16px;
            border-radius: 6px;
            overflow-x: auto;
        }
        a {
            color: #0366d6;
            text-decoration: none;
        }
        a:hover {
            text-decoration: underline;
        }
        .toc {
            background-color: #f9f9f9;
            padding: 1em;
            border-radius: 6px;
            margin-bottom: 2em;
        }
        .toc ul {
            margin: 0;
            padding-left: 1.5em;
        }
    </style>
    """

    MERMAID_SCRIPT = """
    <script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
    <script>
        mermaid.initialize({ startOnLoad: true, theme: 'default' });
    </script>
    """

    def __init__(self, parser: Optional[MarkdownParser] = None):
        self.parser = parser or MarkdownParser()

    def render(
        self,
        markdown_text: str,
        title: str = "Document",
        include_toc: bool = False,
        include_css: bool = True,
        include_mermaid_js: bool = True,
    ) -> str:
        """Render Markdown to complete HTML document."""
        content = self.parser.parse(markdown_text)

        toc_html = ""
        if include_toc:
            toc = self.parser.get_toc()
            if toc:
                toc_html = f'<div class="toc"><strong>Table of Contents</strong>{toc}</div>'

        css = self.DEFAULT_CSS if include_css else ""
        mermaid_js = self.MERMAID_SCRIPT if include_mermaid_js else ""

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    {css}
</head>
<body>
    {toc_html}
    {content}
    {mermaid_js}
</body>
</html>"""

    def render_fragment(self, markdown_text: str) -> str:
        """Render Markdown to HTML fragment (no document wrapper)."""
        return self.parser.parse(markdown_text)


class PDFRenderer:
    """Render Markdown to PDF using WeasyPrint."""

    def __init__(self, parser: Optional[MarkdownParser] = None):
        self.parser = parser or MarkdownParser()
        self.html_renderer = HTMLRenderer(parser=self.parser)

    def render(
        self,
        markdown_text: str,
        title: str = "Document",
        include_toc: bool = False,
    ) -> bytes:
        """Render Markdown to PDF bytes."""
        from weasyprint import HTML, CSS

        # Generate HTML with mermaid disabled (use server-side rendering)
        html_content = self.html_renderer.render(
            markdown_text,
            title=title,
            include_toc=include_toc,
            include_mermaid_js=False,  # WeasyPrint can't execute JS
        )

        # Add print-specific CSS
        print_css = CSS(
            string="""
            @page {
                size: A4;
                margin: 2cm;
            }
            body {
                font-size: 12pt;
            }
        """
        )

        html = HTML(string=html_content)
        return html.write_pdf(stylesheets=[print_css])

    def render_to_file(
        self,
        markdown_text: str,
        output_path: str,
        title: str = "Document",
        include_toc: bool = False,
    ) -> None:
        """Render Markdown to PDF file."""
        pdf_bytes = self.render(markdown_text, title=title, include_toc=include_toc)
        Path(output_path).write_bytes(pdf_bytes)


class WordRenderer:
    """Render Markdown to Word (DOCX) document."""

    def __init__(self, parser: Optional[MarkdownParser] = None):
        self.parser = parser or MarkdownParser()

    def render(
        self,
        markdown_text: str,
        title: str = "Document",
    ) -> bytes:
        """Render Markdown to DOCX bytes."""
        doc = Document()

        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Parse markdown and convert to docx
        self._convert_markdown_to_docx(markdown_text, doc)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def render_to_file(
        self,
        markdown_text: str,
        output_path: str,
        title: str = "Document",
    ) -> None:
        """Render Markdown to DOCX file."""
        docx_bytes = self.render(markdown_text, title=title)
        Path(output_path).write_bytes(docx_bytes)

    def _convert_markdown_to_docx(self, markdown_text: str, doc: Document) -> None:
        """Convert Markdown content to Word document elements."""
        lines = markdown_text.split("\n")
        in_code_block = False
        code_block_content = []
        in_table = False
        table_rows = []

        i = 0
        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.strip().startswith("```"):
                if in_code_block:
                    # End of code block
                    code_text = "\n".join(code_block_content)
                    self._add_code_block(doc, code_text)
                    code_block_content = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue

            # Handle tables
            if "|" in line and not in_table:
                in_table = True
                table_rows = []

            if in_table:
                if "|" in line:
                    # Skip separator row
                    if not re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                        cells = [cell.strip() for cell in line.split("|")[1:-1]]
                        if cells:
                            table_rows.append(cells)
                else:
                    # End of table
                    if table_rows:
                        self._add_table(doc, table_rows)
                    in_table = False
                    table_rows = []
                    continue
                i += 1
                continue

            # Handle headings
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                doc.add_heading(text, level=level)
                i += 1
                continue

            # Handle horizontal rules
            if re.match(r"^[-*_]{3,}$", line.strip()):
                doc.add_paragraph("─" * 50)
                i += 1
                continue

            # Handle blockquotes
            if line.strip().startswith(">"):
                text = line.strip()[1:].strip()
                para = doc.add_paragraph()
                para.style = "Quote"
                para.add_run(text)
                i += 1
                continue

            # Handle lists
            list_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
            if list_match:
                text = list_match.group(2)
                doc.add_paragraph(text, style="List Bullet")
                i += 1
                continue

            ordered_match = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
            if ordered_match:
                text = ordered_match.group(2)
                para = doc.add_paragraph(text, style="List Number")
                i += 1
                continue

            # Handle regular paragraphs
            if line.strip():
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)

            i += 1

        # Handle any remaining table
        if in_table and table_rows:
            self._add_table(doc, table_rows)

    def _add_code_block(self, doc: Document, code: str) -> None:
        """Add a code block to the document."""
        para = doc.add_paragraph()
        run = para.add_run(code)
        run.font.name = "Consolas"
        run.font.size = Pt(10)

    def _add_table(self, doc: Document, rows: list) -> None:
        """Add a table to the document."""
        if not rows:
            return

        num_cols = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"

        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                if j < num_cols:
                    table.rows[i].cells[j].text = cell
                    # Bold header row
                    if i == 0:
                        for para in table.rows[i].cells[j].paragraphs:
                            for run in para.runs:
                                run.bold = True

    def _add_formatted_text(self, para, text: str) -> None:
        """Add text with inline formatting to paragraph."""
        # Simple inline formatting support
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)

        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = para.add_run(part[1:-1])
                run.font.name = "Consolas"
            else:
                para.add_run(part)


class ExcelRenderer:
    """Render Markdown tables to Excel (XLSX) format."""

    def __init__(self, parser: Optional[MarkdownParser] = None):
        self.parser = parser or MarkdownParser()

    def render(
        self,
        markdown_text: str,
        title: str = "Document",
    ) -> bytes:
        """Render Markdown tables to XLSX bytes."""
        wb = Workbook()
        ws = wb.active
        ws.title = title[:31]  # Excel sheet name limit

        tables = self._extract_tables(markdown_text)

        if not tables:
            # If no tables, put all content in first column
            ws.cell(row=1, column=1, value=title)
            ws.cell(row=1, column=1).font = Font(bold=True, size=14)

            lines = markdown_text.strip().split("\n")
            for i, line in enumerate(lines, start=3):
                ws.cell(row=i, column=1, value=line)
        else:
            current_row = 1
            for table_idx, table in enumerate(tables):
                if table_idx > 0:
                    current_row += 2  # Space between tables

                for row_idx, row in enumerate(table):
                    for col_idx, cell in enumerate(row, start=1):
                        excel_cell = ws.cell(
                            row=current_row + row_idx, column=col_idx, value=cell
                        )
                        # Format header row
                        if row_idx == 0:
                            excel_cell.font = Font(bold=True)
                            excel_cell.alignment = Alignment(horizontal="center")

                current_row += len(table)

        # Auto-fit column widths (approximate)
        for column_cells in ws.columns:
            max_length = 0
            column = column_cells[0].column_letter
            for cell in column_cells:
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except Exception:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column].width = adjusted_width

        # Save to bytes
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def render_to_file(
        self,
        markdown_text: str,
        output_path: str,
        title: str = "Document",
    ) -> None:
        """Render Markdown tables to XLSX file."""
        xlsx_bytes = self.render(markdown_text, title=title)
        Path(output_path).write_bytes(xlsx_bytes)

    def _extract_tables(self, markdown_text: str) -> list:
        """Extract tables from Markdown text."""
        tables = []
        lines = markdown_text.split("\n")

        current_table = []
        in_table = False

        for line in lines:
            if "|" in line:
                if not in_table:
                    in_table = True
                    current_table = []

                # Skip separator row
                if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                    continue

                cells = [cell.strip() for cell in line.split("|")[1:-1]]
                if cells:
                    current_table.append(cells)
            else:
                if in_table and current_table:
                    tables.append(current_table)
                in_table = False
                current_table = []

        # Don't forget last table
        if in_table and current_table:
            tables.append(current_table)

        return tables
